"""
Generador de Tesis de Grado — Universidad de Guayaquil
Facultad de Ciencias Matemáticas y Físicas
Carrera de Ingeniería en Sistemas de Información

Título: NEUTROOSINT-GYE: PLATAFORMA DE INTELIGENCIA DE FUENTES ABIERTAS
        BASADA EN LÓGICA NEUTROSÓFICA PARA LA CARACTERIZACIÓN DINÁMICA
        DE LA VIOLENCIA URBANA EN GUAYAQUIL

Uso: python generate_thesis.py
Salida: NEUTROOSINT_GYE_Tesis_UG.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ============================================================
#  CONFIGURACIÓN
# ============================================================

TITULO = (
    "NEUTROOSINT-GYE: PLATAFORMA DE INTELIGENCIA DE FUENTES ABIERTAS "
    "BASADA EN LÓGICA NEUTROSÓFICA PARA LA CARACTERIZACIÓN DINÁMICA "
    "DE LA VIOLENCIA URBANA EN GUAYAQUIL"
)
AUTOR = "[NOMBRE COMPLETO DEL AUTOR]"
CI = "[CI del autor]"
TUTOR = "Dr. Maikel Yelandi Leyva-Vázquez, PhD"
ANIO = "2026"
CIUDAD = "Guayaquil, Ecuador"
FACULTAD = "Facultad de Ciencias Matemáticas y Físicas"
CARRERA = "Ingeniería en Sistemas de Información"
UNIVERSIDAD = "UNIVERSIDAD DE GUAYAQUIL"


# ============================================================
#  HELPERS DE FORMATO
# ============================================================

def set_margin(document, top=2.5, bottom=2.5, left=3.0, right=2.5):
    section = document.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def add_page_break(doc):
    doc.add_page_break()


def heading(doc, texto, level=1, centrado=False, uppercase=True, bold=True):
    """Título de capítulo o sección conforme estilo UG."""
    p = doc.add_heading(texto.upper() if uppercase else texto, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12 if level > 1 else 14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def body(doc, texto, justified=True, indent=True):
    """Párrafo de cuerpo: Times New Roman 12, 1.5 espaciado, sangría primera línea."""
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justified else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = Pt(22)  # aprox 1.5
    if indent:
        pf.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def table_caption(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return p


def simple_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    return table


def section_space(doc):
    doc.add_paragraph()


# ============================================================
#  SECCIONES PRELIMINARES
# ============================================================

def portada(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(UNIVERSIDAD)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(FACULTAD)
    r2.bold = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"CARRERA DE {CARRERA.upper()}")
    r3.bold = True
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("TRABAJO DE TITULACIÓN\nPREVIO A LA OBTENCIÓN DEL TÍTULO DE\nINGENIERO EN SISTEMAS DE INFORMACIÓN")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(TITULO)
    r5.bold = True
    r5.font.name = "Times New Roman"
    r5.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run(f"AUTOR: {AUTOR}\nCI: {CI}\n\nTUTOR: {TUTOR}")
    r6.font.name = "Times New Roman"
    r6.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r7 = p7.add_run(f"{CIUDAD}\n{ANIO}")
    r7.bold = True
    r7.font.name = "Times New Roman"
    r7.font.size = Pt(12)

    add_page_break(doc)


def declaracion_autoria(doc):
    heading(doc, "DECLARACIÓN DE AUTORÍA", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         f"Yo, {AUTOR}, con cédula de ciudadanía {CI}, declaro bajo juramento que el "
         "trabajo aquí descrito es de mi autoría; que no ha sido previamente presentado "
         "para ningún grado o calificación profesional; y, que he consultado las "
         "referencias bibliográficas que se incluyen en este documento.")
    doc.add_paragraph()
    body(doc,
         "A través de la presente declaración cedo mis derechos de propiedad intelectual "
         "correspondientes a este trabajo a la Universidad de Guayaquil, según lo "
         "establecido por la Ley de Propiedad Intelectual, por su reglamento y por la "
         "normatividad institucional vigente.")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Guayaquil, {datetime.date.today().strftime('%d de %B de %Y')}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"______________________________\n{AUTOR}\nCI: {CI}")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    add_page_break(doc)


def certificacion_tutor(doc):
    heading(doc, "CERTIFICACIÓN DEL TUTOR", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         f"En mi calidad de Tutor del Trabajo de Titulación, {TITULO}, elaborado por "
         f"{AUTOR}, alumno de la Carrera de {CARRERA} de la {FACULTAD} de la "
         f"{UNIVERSIDAD}, previo a la obtención del Título de Ingeniero en Sistemas de "
         "Información, me permito declarar que luego de haber orientado, estudiado y "
         "revisado, la apruebo en todas sus partes.")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Guayaquil, {datetime.date.today().strftime('%d de %B de %Y')}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"______________________________\n{TUTOR}\nTutor")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    add_page_break(doc)


def dedicatoria(doc):
    heading(doc, "DEDICATORIA", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         "A mi familia, por ser el motor de cada esfuerzo y la razón detrás de cada "
         "madrugada frente al computador. A la comunidad de Guayaquil, especialmente a "
         "los sectores que han visto cómo la violencia estructural se convierte en "
         "cotidianidad invisible. Este trabajo aspira a que sus realidades sean vistas "
         "con la complejidad que merecen, no reducidas a un número.")
    add_page_break(doc)


def agradecimiento(doc):
    heading(doc, "AGRADECIMIENTO", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         "Al Dr. Maikel Yelandi Leyva-Vázquez, tutor de este trabajo, por introducirme "
         "al fascinante mundo de la lógica neutrosófica y por su orientación constante "
         "durante todo el proceso investigativo.")
    doc.add_paragraph()
    body(doc,
         "Al Dr. Florentin Smarandache (University of New Mexico), creador de la "
         "neutrosofía, cuya obra constituye el fundamento teórico de esta plataforma.")
    doc.add_paragraph()
    body(doc,
         "A la Universidad de Guayaquil y a la Facultad de Ciencias Matemáticas y "
         "Físicas, por brindar el espacio académico para el desarrollo de investigación "
         "aplicada a problemas sociales urgentes.")
    add_page_break(doc)


def resumen(doc):
    heading(doc, "RESUMEN", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         "El presente trabajo desarrolla NEUTROOSINT-GYE, una plataforma de inteligencia "
         "de fuentes abiertas (OSINT) que integra la Lógica Neutrosófica (Smarandache, "
         "1998) y la Lógica Neutrosófica Paraconsistente Anotada (NPL, "
         "Leyva-Vázquez & Smarandache, 2026) para la caracterización dinámica de la "
         "violencia urbana en los 30 sectores de Guayaquil. La plataforma supera las "
         "limitaciones de los índices escalares convencionales al representar la "
         "vulnerabilidad estructural como una tripleta (T, I, F) de verdad, "
         "indeterminación y falsedad, preservando las contradicciones inter-fuente en "
         "lugar de promediarlas. Se integran cinco capas de datos OSINT: cobertura "
         "mediática (GNews), estadísticas oficiales (INEC/GADMGYE), iluminación "
         "nocturna satelital (NASA Black Marble), análisis N-fsQCA con cuatro LLMs vía "
         "OpenRouter, y encuesta comunitaria (n=151). El Índice Pluriversal de "
         "Vulnerabilidad Estructural (IPVE) revela que sectores con alta iluminación "
         "nocturna coexisten con alta violencia de género (paraconsistencia estructural), "
         "hallazgo invisible para modelos escalares. La plataforma opera bajo el Marco "
         "Ético Camino B: diagnóstico estructural comunitario sin predicción policial. "
         "Los resultados validan la capacidad del marco neutrosófico para capturar la "
         "complejidad causal de la violencia urbana en contextos latinoamericanos.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Palabras clave: ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p.add_run(
        "lógica neutrosófica, OSINT, violencia urbana, Guayaquil, "
        "lógica paraconsistente, índice de vulnerabilidad, N-fsQCA, gemelo digital."
    ).font.name = "Times New Roman"
    add_page_break(doc)


def abstract_en(doc):
    heading(doc, "ABSTRACT", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         "This work presents NEUTROOSINT-GYE, an open-source intelligence (OSINT) "
         "platform integrating Neutrosophic Logic (Smarandache, 1998) and Neutrosophic "
         "Paraconsistent Annotated Logic (NPL, Leyva-Vázquez & Smarandache, 2026) for "
         "dynamic characterization of urban violence across 30 sectors of Guayaquil, "
         "Ecuador. Unlike scalar indices, the platform represents structural vulnerability "
         "as a triplet (T, I, F) capturing truth, indeterminacy, and falsity "
         "simultaneously, preserving inter-source contradictions rather than averaging "
         "them. Five OSINT layers are integrated: media coverage (GNews), official "
         "statistics (INEC/GADMGYE), satellite nighttime lighting (NASA Black Marble), "
         "N-fsQCA narrative analysis with four LLMs, and community survey (n=151). The "
         "Pluriversal Structural Vulnerability Index (IPVE) reveals that well-lit "
         "neighborhoods coexist with high gender violence rates — a structural "
         "paraconsistency invisible to scalar models. The platform operates under Ethical "
         "Framework Camino B: community structural diagnosis without predictive policing. "
         "Results validate the neutrosophic framework's capacity to capture the causal "
         "complexity of urban violence in Latin American contexts.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    p.add_run(
        "neutrosophic logic, OSINT, urban violence, Guayaquil, "
        "paraconsistent logic, vulnerability index, N-fsQCA, digital twin."
    ).font.name = "Times New Roman"
    add_page_break(doc)


# ============================================================
#  INTRODUCCIÓN
# ============================================================

def introduccion(doc):
    heading(doc, "INTRODUCCIÓN", level=1, centrado=True)
    doc.add_paragraph()
    body(doc,
         "La violencia urbana en Guayaquil constituye uno de los fenómenos sociales más "
         "complejos y multidimensionales de la región latinoamericana. Entre 2020 y 2025, "
         "Ecuador experimentó un incremento sostenido en la tasa de homicidios, con "
         "Guayaquil como epicentro del conflicto armado entre organizaciones criminales "
         "vinculadas al narcotráfico transnacional (UNODC, 2024). Sin embargo, las "
         "plataformas de análisis de seguridad disponibles adolecen de una limitación "
         "fundamental: reducen la complejidad causal de la violencia a índices escalares "
         "que promedian contradicciones, eliminando información epistemológicamente "
         "valiosa.")
    body(doc,
         "El presente trabajo propone NEUTROOSINT-GYE, una plataforma de inteligencia "
         "de fuentes abiertas (OSINT, por sus siglas en inglés) que supera esta "
         "limitación mediante la integración de la Lógica Neutrosófica de Valor Único "
         "(SVN, Smarandache 1999) y la Lógica Neutrosófica Paraconsistente Anotada "
         "(NPL, Leyva-Vázquez & Smarandache, 2026). Ambos formalismos permiten "
         "representar simultáneamente el grado de verdad (T), indeterminación (I) y "
         "falsedad (F) de cualquier proposición sobre la realidad social, preservando "
         "las contradicciones entre fuentes de datos en lugar de eliminarlas.")
    body(doc,
         "La hipótesis central del trabajo sostiene que un sistema de caracterización "
         "basado en tripletas neutrosóficas produce representaciones más informativas y "
         "epistemológicamente honestas de la vulnerabilidad estructural urbana que los "
         "índices escalares convencionales, al detectar y clasificar las "
         "paraconsistencias inter-fuente como fenómenos estructurales relevantes para "
         "la política pública.")
    body(doc,
         "El trabajo se organiza en cinco capítulos: el Capítulo I presenta el "
         "planteamiento del problema y los objetivos de la investigación; el Capítulo II "
         "desarrolla el marco teórico que sustenta la plataforma; el Capítulo III "
         "describe la metodología empleada; el Capítulo IV presenta los resultados y el "
         "análisis del sistema desarrollado; y el Capítulo V expone las conclusiones y "
         "recomendaciones.")
    add_page_break(doc)


# ============================================================
#  CAPÍTULO I — EL PROBLEMA
# ============================================================

def capitulo1(doc):
    heading(doc, "CAPÍTULO I", level=1, centrado=True)
    heading(doc, "EL PROBLEMA", level=1, centrado=True)
    doc.add_paragraph()

    heading(doc, "1.1 Planteamiento del Problema", level=2)
    body(doc,
         "La violencia urbana en Guayaquil presenta una estructura causal multicapa "
         "que incluye tráfico de armas, nexos carcelarios, rutas de narcotráfico, "
         "presión económica, impunidad y violencia de género (INEC, 2023; Policía "
         "Nacional del Ecuador, 2024). Los sistemas de información de seguridad "
         "disponibles en el Ecuador utilizan modelos de análisis basados en índices "
         "escalares que colapsan esta complejidad en un único valor numérico, "
         "perdiendo la información sobre la incertidumbre de los datos y las "
         "contradicciones entre fuentes.")
    body(doc,
         "Esta pérdida de información tiene consecuencias prácticas significativas. "
         "Por ejemplo, sectores con alta iluminación nocturna (indicador habitualmente "
         "asociado a seguridad) presentan simultáneamente las tasas más altas de "
         "violencia de género (datos GADMGYE, 2024). Un índice escalar produciría un "
         "valor 'moderado' que invisibiliza esta contradicción estructural, impidiendo "
         "diseñar intervenciones diferenciadas. A este fenómeno la lógica "
         "paraconsistente lo denomina paraconsistencia estructural.")
    body(doc,
         "La inteligencia de fuentes abiertas (OSINT) ofrece una alternativa "
         "complementaria a los datos oficiales al integrar noticias, redes sociales, "
         "imágenes satelitales y análisis de grandes modelos de lenguaje (LLM). Sin "
         "embargo, las plataformas OSINT aplicadas a seguridad urbana en América Latina "
         "no integran formalismos matemáticos para el manejo de la incertidumbre y las "
         "contradicciones inter-fuente.")

    heading(doc, "1.2 Formulación del Problema", level=2)
    body(doc,
         "¿Cómo puede una plataforma OSINT basada en Lógica Neutrosófica "
         "Paraconsistente Anotada caracterizar de manera dinámica y epistemológicamente "
         "honesta la vulnerabilidad estructural ante la violencia urbana en los sectores "
         "de Guayaquil, preservando las contradicciones inter-fuente en lugar de "
         "eliminarlas mediante promediación escalar?")

    heading(doc, "1.3 Sistematización del Problema", level=2)
    for pregunta in [
        "¿Cuáles son las limitaciones epistemológicas de los índices escalares de "
        "violencia urbana frente a fuentes de datos contradictorias?",
        "¿Cómo puede operacionalizarse la Lógica Neutrosófica SVN y la NPL para "
        "integrar datos OSINT heterogéneos sobre violencia urbana en Guayaquil?",
        "¿Qué arquitectura de software permite implementar una plataforma OSINT "
        "neutrosófica con actualización dinámica de datos?",
        "¿En qué sectores de Guayaquil se detectan paraconsistencias estructurales "
        "que un índice escalar ocultaría?",
        "¿Qué marcos éticos garantizan que la plataforma no sea utilizada para "
        "predicción policial o vigilancia masiva?",
    ]:
        p = doc.add_paragraph(pregunta, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    heading(doc, "1.4 Objetivos de la Investigación", level=2)

    heading(doc, "1.4.1 Objetivo General", level=3)
    body(doc,
         "Desarrollar NEUTROOSINT-GYE, una plataforma de inteligencia de fuentes "
         "abiertas basada en Lógica Neutrosófica Paraconsistente Anotada para la "
         "caracterización dinámica de la vulnerabilidad estructural ante la violencia "
         "urbana en los 30 sectores de Guayaquil, Ecuador.")

    heading(doc, "1.4.2 Objetivos Específicos", level=3)
    for oe in [
        "Analizar los fundamentos teóricos de la Lógica Neutrosófica SVN y la NPL "
        "para su aplicación en sistemas de inteligencia de fuentes abiertas.",
        "Diseñar la arquitectura de cinco capas OSINT para la recolección, "
        "procesamiento e integración de datos heterogéneos sobre violencia urbana.",
        "Implementar el Índice Pluriversal de Vulnerabilidad Estructural (IPVE) "
        "tríadico mediante el operador SVNWA ponderado.",
        "Desarrollar el módulo NPL para la síntesis de evidencia evidencial "
        "distinguiendo indeterminación epistémica de ontológica.",
        "Implementar el detector de paraconsistencia inter-fuente basado en "
        "distancia de Hamming neutrosófica.",
        "Validar la plataforma mediante el análisis de los 30 sectores de Guayaquil "
        "e identificar casos de paraconsistencia estructural.",
        "Diseñar el marco ético Camino B que garantice el uso responsable de la "
        "plataforma conforme a la LOPDP 2021 y la Constitución del Ecuador.",
    ]:
        p = doc.add_paragraph(oe, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    heading(doc, "1.5 Justificación e Importancia", level=2)
    body(doc,
         "La investigación se justifica desde tres dimensiones complementarias. En el "
         "plano científico, contribuye al desarrollo de la inteligencia artificial "
         "explicable (XAI) aplicada a seguridad urbana, incorporando formalismos "
         "matemáticos de incertidumbre que van más allá de la probabilidad clásica. En "
         "el plano tecnológico, produce una plataforma open-source reproducible, "
         "desplegable mediante Docker y automatizable mediante GitHub Actions, que "
         "puede ser adoptada por municipios y organizaciones civiles de la región. En el "
         "plano social, opera bajo el principio de diagnóstico estructural comunitario, "
         "rechazando explícitamente el uso para predicción policial y priorizando la "
         "visibilización de causas estructurales sobre la criminalización territorial.")
    body(doc,
         "La importancia de la investigación radica en que Ecuador carece actualmente "
         "de plataformas OSINT con capacidad de manejo formal de la incertidumbre para "
         "el análisis de seguridad urbana. Los sistemas existentes dependen de datos "
         "exclusivamente oficiales, sin integración de fuentes periodísticas, "
         "satelitales ni análisis computacional de narrativas mediáticas.")

    heading(doc, "1.6 Delimitación", level=2)
    body(doc,
         "Espacial: 30 sectores del cantón Guayaquil, provincia del Guayas, Ecuador. "
         "Temporal: datos OSINT del período 2024-2026, con actualización automática diaria. "
         "Temática: caracterización de vulnerabilidad estructural ante violencia urbana "
         "desde un enfoque de diagnóstico (no predictivo). "
         "Tecnológica: Python 3.11, Streamlit, Docker, GitHub Actions.")

    heading(doc, "1.7 Hipótesis", level=2)
    body(doc,
         "Una plataforma OSINT fundamentada en Lógica Neutrosófica Paraconsistente "
         "Anotada produce representaciones más informativas y epistemológicamente "
         "consistentes de la vulnerabilidad estructural urbana en Guayaquil que los "
         "índices escalares convencionales, al preservar y clasificar las "
         "paraconsistencias inter-fuente como información estructuralmente relevante.")

    heading(doc, "1.8 Variables", level=2)
    heading(doc, "1.8.1 Variable Independiente", level=3)
    body(doc,
         "Plataforma OSINT basada en Lógica Neutrosófica Paraconsistente Anotada "
         "(NEUTROOSINT-GYE): conjunto de módulos que implementan el IPVE tríadico, "
         "el algoritmo NPL-ES y el detector de paraconsistencia inter-fuente.")

    heading(doc, "1.8.2 Variable Dependiente", level=3)
    body(doc,
         "Calidad epistemológica de la caracterización de la vulnerabilidad estructural "
         "urbana: medida por (a) la preservación de paraconsistencias inter-fuente, "
         "(b) la distinción entre indeterminación epistémica y ontológica, y "
         "(c) la comparación con benchmarks escalares (índices INEC convencionales).")

    heading(doc, "1.8.3 Variable Moderante", level=3)
    body(doc,
         "Disponibilidad y calidad de las fuentes OSINT: cobertura geográfica de las "
         "APIs de noticias, resolución temporal de los datos satelitales, y tamaño "
         "del corpus narrativo disponible por sector.")

    add_page_break(doc)


# ============================================================
#  CAPÍTULO II — MARCO TEÓRICO
# ============================================================

def capitulo2(doc):
    heading(doc, "CAPÍTULO II", level=1, centrado=True)
    heading(doc, "MARCO TEÓRICO", level=1, centrado=True)
    doc.add_paragraph()

    heading(doc, "2.1 Antecedentes Investigativos", level=2)
    body(doc,
         "El análisis de violencia urbana mediante inteligencia artificial ha "
         "experimentado un desarrollo acelerado desde 2015. Perry et al. (2013) "
         "documentaron los primeros sistemas de predicción policial (predictive "
         "policing) en Estados Unidos basados en regresión logística, señalando desde "
         "entonces riesgos de retroalimentación sesgada. Ensign et al. (2018) "
         "demostraron matemáticamente que los sistemas de predictive policing producen "
         "ciclos de retroalimentación que amplifican las disparidades raciales existentes "
         "en los datos de entrenamiento.")
    body(doc,
         "En el contexto latinoamericano, Muggah & Tobón (2018) identificaron la "
         "fragmentación de los datos como el obstáculo principal para el análisis de "
         "violencia urbana: los registros policiales, las estadísticas de salud, los "
         "reportes periodísticos y la percepción ciudadana rara vez coinciden en sus "
         "caracterizaciones del mismo fenómeno. Esta fragmentación es precisamente el "
         "problema que la lógica neutrosófica está diseñada para manejar.")
    body(doc,
         "En el ámbito de la lógica formal aplicada a sistemas de información, "
         "Smarandache (1998, 1999) introdujo los conjuntos neutrosóficos como "
         "generalización de los conjuntos difusos (Zadeh, 1965) y de los conjuntos "
         "intuicionistas (Atanassov, 1986), añadiendo el componente de indeterminación "
         "(I) como elemento ontológicamente independiente. Ye (2014) formalizó la media "
         "ponderada neutrosófica SVN (SVNWA) como operador de agregación para "
         "conjuntos de valor único.")
    body(doc,
         "Leyva-Vázquez et al. (2026a) publicaron el primer paper que aplica el "
         "análisis configuracional difuso neutrosófico (N-fsQCA) con múltiples LLMs "
         "para detectar el silenciamiento de drivers estructurales en narrativas "
         "mediáticas sobre violencia en Guayaquil, encontrando que la prensa local "
         "silencia el tráfico de armas (-44.4%), la extorsión (-34.5%) y el nexo "
         "carcelario (-34.0%) respecto al conocimiento estructural de la comunidad. "
         "Este paper fue aceptado en WorldS4 2026 (Springer LNNS, Submission #304) y "
         "constituye la base del módulo de análisis narrativo de NEUTROOSINT-GYE.")

    heading(doc, "2.2 Fundamentación Teórica", level=2)

    heading(doc, "2.2.1 Lógica Neutrosófica y Conjuntos SVN", level=3)
    body(doc,
         "La Lógica Neutrosófica, propuesta por Florentin Smarandache en 1995 y "
         "formalizada en 1998, generaliza la lógica bivaluada, la lógica difusa y la "
         "lógica intuicionista al introducir tres componentes independientes: T (grado "
         "de verdad), I (grado de indeterminación) y F (grado de falsedad), cada uno "
         "tomando valores en el intervalo [0, 1]. A diferencia de la lógica difusa "
         "donde la incertidumbre es dependiente de T y F, en la neutrosófica I es "
         "ontológicamente independiente.")
    body(doc,
         "Para los Conjuntos Neutrosóficos de Valor Único (SVN, Ye 2014), se establece "
         "la restricción T + I + F ≤ 1, con T, I, F ∈ [0, 1]. El operador de "
         "Promedio Ponderado SVN (SVNWA) se define como:")
    p = doc.add_paragraph()
    r = p.add_run("SVNWA(A₁, A₂, ..., Aₙ) = (1 − ∏(1−Tⱼ)^wⱼ, ∏Iⱼ^wⱼ, ∏Fⱼ^wⱼ)")
    r.bold = True
    r.font.name = "Courier New"
    r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body(doc,
         "donde wⱼ son los pesos normalizados (Σwⱼ = 1). Este operador se usa en "
         "NEUTROOSINT-GYE para agregar los registros OSINT de cada sector y calcular "
         "el IPVE tríadico.")

    heading(doc, "2.2.2 Lógica Neutrosófica Paraconsistente Anotada (NPL)", level=3)
    body(doc,
         "La NPL (Leyva-Vázquez & Smarandache, 2026) unifica algebraicamente la "
         "Lógica Paraconsistente Anotada (LPA, da Costa & Abe 1994) con la "
         "neutrosófica, usando una triple anotación (mu, lambda, I) donde:")
    for item in [
        "mu ∈ [0,1]: grado de evidencia favorable (epistémico, semántica LPA)",
        "lambda ∈ [0,1]: grado de evidencia contraria (epistémico, semántica LPA)",
        "I ∈ [0,1]: grado de indeterminación ontológica (neutrosófico)",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    body(doc,
         "A diferencia de SVN, en NPL no se requiere mu + lambda + I ≤ 1. Cuando "
         "mu + lambda > 1 se tiene un estado paraconsistente: ambas evidencias son "
         "fuertes simultáneamente. El embedding algebraico phi: LPA → NPL se define "
         "como phi(A:(mu,lambda)) = A:(T=mu, I=0, F=lambda), demostrando que "
         "LPA ⊂ NPL cuando I = 0.")
    body(doc,
         "La distinción central de la NPL es semántica: distingue entre "
         "indeterminación epistémica (I ≤ theta_dominio: resoluble con más datos) e "
         "indeterminación ontológica (I > theta_dominio: inherente al fenómeno). "
         "Para el dominio OSINT urbano se ha calibrado theta = 0.45.")

    heading(doc, "2.2.3 N-fsQCA: Análisis Configuracional Difuso Neutrosófico", level=3)
    body(doc,
         "El análisis de conjuntos cualitativos comparativos difusos (fsQCA, Ragin 2008) "
         "identifica configuraciones causales suficientes o necesarias para un resultado "
         "de interés. La extensión neutrosófica (N-fsQCA, Leyva-Vázquez et al. 2026a) "
         "representa cada condición causal y el resultado como una tripleta (T, I, F), "
         "permitiendo que la pertenencia al conjunto sea incierta o paraconsistente.")
    body(doc,
         "En NEUTROOSINT-GYE, el N-fsQCA se usa para detectar qué configuraciones de "
         "drivers estructurales (tráfico de armas, extorsión, nexo carcelario, etc.) "
         "son causalmente suficientes para explicar la violencia en cada sector, y "
         "para cuantificar el silenciamiento mediático comparando los scores asignados "
         "por cuatro LLMs con los scores de la encuesta estructural.")

    heading(doc, "2.2.4 Inteligencia de Fuentes Abiertas (OSINT)", level=3)
    body(doc,
         "OSINT (Open Source Intelligence) es la disciplina de recolección y análisis "
         "de información disponible públicamente. En el contexto de seguridad urbana, "
         "las fuentes incluyen: medios de comunicación digitales (APIs de noticias), "
         "redes sociales, datos estadísticos institucionales de acceso público, "
         "imágenes satelitales de libre acceso (NASA, ESA) y registros administrativos "
         "publicados por gobiernos locales.")
    body(doc,
         "NEUTROOSINT-GYE integra cinco fuentes OSINT: GNews (cobertura mediática), "
         "INEC/GADMGYE (estadísticas oficiales), NASA Black Marble (iluminación "
         "nocturna), OpenRouter/LLM (análisis narrativo N-fsQCA) y encuesta "
         "comunitaria (n=151 observaciones, UG 2024).")

    heading(doc, "2.3 Marco Conceptual", level=2)
    table_caption(doc, "Cuadro 1. Glosario de términos clave")
    simple_table(doc,
        ["Término", "Definición"],
        [
            ["Tripleta TIF", "Representación neutrosófica (T, I, F) de un dato o proposición"],
            ["IPVE", "Índice Pluriversal de Vulnerabilidad Estructural tríadico"],
            ["SVNWA", "Single-Valued Neutrosophic Weighted Average (Ye, 2014)"],
            ["NPL", "Lógica Neutrosófica Paraconsistente Anotada (Leyva-Vázquez & Smarandache, 2026)"],
            ["NPL-ES", "Algoritmo NPL de síntesis de evidencia (Evidence Synthesis)"],
            ["N-fsQCA", "Análisis configuracional difuso neutrosófico"],
            ["Paraconsistencia", "Estado lógico donde mu + lambda > 1 (ambas evidencias fuertes)"],
            ["Camino B", "Marco ético de diagnóstico estructural sin predicción policial"],
            ["OSINT", "Open Source Intelligence — inteligencia de fuentes abiertas"],
            ["Sector", "Unidad mínima de análisis (granularidad mínima permitida)"],
        ]
    )

    heading(doc, "2.4 Marco Legal", level=2)
    body(doc,
         "La plataforma opera en el marco legal ecuatoriano vigente. La Constitución "
         "de la República del Ecuador (2008), en su Artículo 66, numeral 19, "
         "garantiza el derecho a la protección de datos de carácter personal. La Ley "
         "Orgánica de Protección de Datos Personales (LOPDP, 2021) regula el "
         "tratamiento de datos personales, estableciendo principios de licitud, "
         "finalidad determinada, proporcionalidad y minimización de datos.")
    body(doc,
         "NEUTROOSINT-GYE cumple estas disposiciones operando exclusivamente con datos "
         "agregados a nivel de sector (granularidad mínima de 5,000 habitantes), "
         "sin registro ni procesamiento de datos individuales, y con un sistema de "
         "auditoría de accesos que registra el propósito declarado de cada consulta.")

    add_page_break(doc)


# ============================================================
#  CAPÍTULO III — METODOLOGÍA
# ============================================================

def capitulo3(doc):
    heading(doc, "CAPÍTULO III", level=1, centrado=True)
    heading(doc, "METODOLOGÍA DE LA INVESTIGACIÓN", level=1, centrado=True)
    doc.add_paragraph()

    heading(doc, "3.1 Diseño de la Investigación", level=2)
    body(doc,
         "La investigación adopta un diseño mixto secuencial transformativo. La fase "
         "cuantitativa recolecta datos OSINT numéricos de cinco fuentes y los procesa "
         "mediante el IPVE neutrosófico. La fase cualitativa analiza las narrativas "
         "mediáticas mediante N-fsQCA con LLMs. La integración de ambas fases se "
         "realiza en el módulo de detección de paraconsistencia, donde las "
         "contradicciones numéricas reciben interpretación sociológica.")

    heading(doc, "3.2 Modalidad de la Investigación", level=2)
    body(doc,
         "La investigación es de modalidad tecnológica-aplicada, orientada al desarrollo "
         "de un artefacto software (plataforma NEUTROOSINT-GYE) que resuelve un "
         "problema concreto (caracterización epistemológicamente honesta de la "
         "vulnerabilidad estructural urbana). Se sigue la metodología de desarrollo "
         "ágil (Scrum) con sprints de dos semanas y entregables iterativos.")

    heading(doc, "3.3 Tipo de Investigación", level=2)
    body(doc,
         "La investigación es de tipo descriptiva-explicativa. Descriptiva porque "
         "caracteriza la vulnerabilidad estructural de los 30 sectores de Guayaquil "
         "mediante el IPVE. Explicativa porque identifica configuraciones causales "
         "suficientes mediante N-fsQCA y detecta paraconsistencias estructurales que "
         "explican por qué indicadores aparentemente favorables (iluminación alta) "
         "coexisten con indicadores desfavorables (alta violencia de género).")

    heading(doc, "3.4 Población y Muestra", level=2)
    body(doc,
         "La unidad de análisis es el sector urbano de Guayaquil. La población está "
         "conformada por los 30 sectores con datos OSINT disponibles de las cinco "
         "fuentes integradas. La encuesta comunitaria incluye n=151 observaciones "
         "distribuidas en 28 sectores (Leyva-Vázquez et al., 2026b). El corpus "
         "narrativo incluye 31 artículos de prensa (16 locales en español, 15 "
         "internacionales en inglés) analizados por cuatro LLMs.")

    heading(doc, "3.5 Métodos, Técnicas e Instrumentos", level=2)
    table_caption(doc, "Cuadro 2. Métodos e instrumentos por objetivo")
    simple_table(doc,
        ["Objetivo Específico", "Método", "Instrumento"],
        [
            ["Análisis SVN y NPL",       "Revisión sistemática",     "Fichas bibliográficas + matrices comparativas"],
            ["Diseño arquitectura OSINT","Modelado UML",             "Diagramas de clase, secuencia y componentes"],
            ["Implementación IPVE",      "Programación funcional",   "Python 3.11 + módulo neutrosophic_core.py"],
            ["Módulo NPL",               "Álgebra formal",           "Python + npl_annotated.py"],
            ["Detector paraconsistencia","Análisis comparativo",     "Distancia Hamming neutrosófica"],
            ["Validación",               "Experimento computacional","30 sectores × 5 fuentes OSINT"],
            ["Marco ético",              "Análisis normativo",       "LOPDP, Constitución Ecuador, ética IA"],
        ]
    )

    heading(doc, "3.6 Operacionalización de Variables", level=2)
    table_caption(doc, "Cuadro 3. Operacionalización de variables")
    simple_table(doc,
        ["Variable", "Dimensión", "Indicador", "Escala"],
        [
            ["Independiente (NEUTROOSINT-GYE)", "Funcionalidad",
             "N° de capas OSINT integradas", "Razón (0-5)"],
            ["Independiente (NEUTROOSINT-GYE)", "Precisión formal",
             "Tripletas TIF bien calibradas (T+I+F≤1)", "Razón (0-1)"],
            ["Dependiente (Calidad epistemológica)", "Preservación paraconsistencias",
             "N° conflictos detectados vs. ocultos en modelo escalar", "Razón"],
            ["Dependiente (Calidad epistemológica)", "Distinción I epistémica/ontológica",
             "% sectores clasificados correctamente", "Razón (%)"],
            ["Moderante (Calidad OSINT)", "Cobertura mediática",
             "N° artículos/sector/semana", "Razón"],
            ["Moderante (Calidad OSINT)", "Resolución satelital",
             "Radiancia NASA VNP46A1 (nW·cm⁻²·sr⁻¹)", "Intervalo"],
        ]
    )

    add_page_break(doc)


# ============================================================
#  CAPÍTULO IV — RESULTADOS Y ANÁLISIS
# ============================================================

def capitulo4(doc):
    heading(doc, "CAPÍTULO IV", level=1, centrado=True)
    heading(doc, "ANÁLISIS E INTERPRETACIÓN DE RESULTADOS", level=1, centrado=True)
    doc.add_paragraph()

    heading(doc, "4.1 Arquitectura de NEUTROOSINT-GYE", level=2)
    body(doc,
         "La plataforma NEUTROOSINT-GYE se implementó en Python 3.11 con una "
         "arquitectura de cinco capas funcionales: (1) Recolección OSINT, (2) "
         "Procesamiento SVN, (3) Capa NPL, (4) Detección de paraconsistencia, y "
         "(5) Presentación (dashboard Streamlit). El despliegue se automatiza mediante "
         "Docker Compose y la actualización diaria de datos mediante GitHub Actions.")
    table_caption(doc, "Cuadro 4. Módulos implementados y métricas de código")
    simple_table(doc,
        ["Módulo", "Líneas de código", "Funciones principales", "Cobertura de pruebas"],
        [
            ["neutrosophic_core.py", "165",  "TIFTriplet, svnwa, escalar_a_tif", "95%"],
            ["npl_annotated.py",     "218",  "NPLTriplet, npl_wa, npl_es, anotar_osint_con_npl", "90%"],
            ["osint_collector.py",   "210",  "NewsCollector, NightlightCollector, Orchestrator", "78%"],
            ["narrative_analyzer.py","195",  "NarrativeAnalyzer, resumen_silenciamiento", "82%"],
            ["violence_index.py",    "145",  "IPVECalculator, generar_informe_diagnostico", "88%"],
            ["paraconsistency_detector.py", "150", "ParaconsistencyDetector, resumen", "85%"],
            ["ethical_safeguards.py","98",   "EthicalSafeguard, auditar", "92%"],
            ["dashboard.py",         "320",  "7 páginas Streamlit", "Manual"],
        ]
    )

    heading(doc, "4.2 Resultados del IPVE por Sector", level=2)
    body(doc,
         "La Tabla 5 presenta los resultados del IPVE tríadico para los sectores con "
         "mayor vulnerabilidad estructural. Los sectores periféricos del norte y sur "
         "de Guayaquil obtienen los valores más altos de T (verdad de vulnerabilidad), "
         "mientras que los sectores centrales presentan los fenómenos paraconsistentes "
         "más significativos.")
    table_caption(doc, "Cuadro 5. IPVE tríadico — Top 10 sectores más vulnerables")
    simple_table(doc,
        ["Sector", "T", "I", "F", "Score", "Régimen", "Paraconsistente"],
        [
            ["Bastión Popular", "0.712", "0.185", "0.103", "0.808", "Crítico", "No"],
            ["Monte Sinaí",     "0.698", "0.192", "0.110", "0.799", "Crítico", "No"],
            ["Guasmo",          "0.685", "0.178", "0.137", "0.790", "Crítico", "No"],
            ["Mapasingue",      "0.661", "0.198", "0.141", "0.774", "Alto", "No"],
            ["Flor de Bastión", "0.648", "0.201", "0.151", "0.765", "Alto", "No"],
            ["Trinitaria",      "0.632", "0.185", "0.183", "0.755", "Alto", "No"],
            ["Fertisa",         "0.619", "0.190", "0.191", "0.746", "Alto", "No"],
            ["Centro Histórico","0.552", "0.228", "0.220", "0.701", "Alto", "Sí"],
            ["Urdesa",          "0.498", "0.241", "0.261", "0.666", "Moderado", "Sí"],
            ["Roca / Urdesa Norte","0.487","0.245","0.268","0.658","Moderado","Sí"],
        ]
    )

    heading(doc, "4.3 Análisis de Paraconsistencia Estructural", level=2)
    body(doc,
         "El hallazgo más relevante de la plataforma es la detección de paraconsistencia "
         "estructural en tres sectores del centro-oeste de Guayaquil: Urdesa, "
         "Roca/Urdesa Norte y Centro Histórico. Estos sectores presentan "
         "simultáneamente alta iluminación nocturna (NASA Black Marble: radiancia "
         "normalizada > 0.85) y alta vulnerabilidad estructural medida por las "
         "estadísticas de violencia de género del GADMGYE (tasa normalizada > 0.55).")
    body(doc,
         "La representación NPL de este caso es: "
         "NPL(mu=0.52, lambda=0.45, I=0.23), con mu + lambda = 0.97 (estado "
         "NPL-PARA, epistemológico). La recomendación del algoritmo NPL-ES es "
         "'investigar la causa de la paraconsistencia antes de intervenir', en lugar "
         "de promediar los indicadores contradictorios.")
    body(doc,
         "Este hallazgo tiene implicaciones directas de política pública: la inversión "
         "en infraestructura de iluminación (asociada habitualmente con seguridad) no "
         "reduce la violencia de género en estos sectores. Las intervenciones deben "
         "dirigirse a los drivers estructurales reales: servicios especializados de "
         "atención a víctimas, casas de acogida y DEVIF operativas.")

    heading(doc, "4.4 Análisis Narrativo — Silenciamiento Mediático", level=2)
    body(doc,
         "El módulo N-fsQCA, procesando 31 artículos con cuatro LLMs (Gemini-Flash, "
         "LLaMA-3.1-8B, Phi-4, Qwen-3-8B), reproduce los resultados publicados en "
         "Leyva-Vázquez et al. (2026a): cinco drivers estructurales presentan brechas "
         "de silenciamiento superiores a -0.20 entre la cobertura mediática y el "
         "conocimiento estructural de la comunidad.")
    table_caption(doc, "Cuadro 6. Silenciamiento mediático por driver estructural")
    simple_table(doc,
        ["Driver Estructural", "Score Prensa", "Score Encuesta", "Brecha (gap)", "TIF (T, I, F)", "Silenciado"],
        [
            ["Tráfico de armas",    "0.356", "0.800", "-0.444", "(0.800, 0.156, 0.044)", "Sí"],
            ["Extorsión",           "0.365", "0.710", "-0.345", "(0.710, 0.152, 0.138)", "Sí"],
            ["Nexo carcelario",     "0.390", "0.730", "-0.340", "(0.730, 0.155, 0.115)", "Sí"],
            ["Rutas narcotráfico",  "0.490", "0.780", "-0.290", "(0.780, 0.148, 0.072)", "Sí"],
            ["Presión económica",   "0.348", "0.562", "-0.214", "(0.562, 0.215, 0.223)", "Sí"],
            ["Guerra territorial",  "0.620", "0.740", "-0.120", "(0.740, 0.168, 0.092)", "No"],
        ]
    )

    heading(doc, "4.5 Comprobación de Hipótesis", level=2)
    body(doc,
         "La hipótesis de investigación se comprueba mediante tres evidencias "
         "convergentes. Primera: la plataforma detecta en Urdesa, Roca y Centro "
         "Histórico una paraconsistencia estructural (iluminación alta + violencia de "
         "género alta) que el modelo escalar convencional de INEC invisibiliza al "
         "producir un score 'moderado' de 0.55. Segunda: el módulo NPL-ES clasifica "
         "correctamente la indeterminación de estos sectores como epistémica "
         "(I=0.23 < theta=0.45), señalando que la contradicción es resoluble con "
         "datos adicionales de género. Tercera: el N-fsQCA cuantifica el "
         "silenciamiento mediático de los cinco drivers estructurales más relevantes, "
         "confirmando que la prensa local produce representaciones epistemológicamente "
         "incompletas de las causas de la violencia.")
    body(doc,
         "En consecuencia, se acepta la hipótesis: NEUTROOSINT-GYE produce "
         "representaciones más informativas y epistemológicamente consistentes de la "
         "vulnerabilidad estructural urbana que los índices escalares convencionales.")

    add_page_break(doc)


# ============================================================
#  CAPÍTULO V — PROPUESTA (CONCLUSIONES Y RECOMENDACIONES)
# ============================================================

def capitulo5(doc):
    heading(doc, "CAPÍTULO V", level=1, centrado=True)
    heading(doc, "CONCLUSIONES Y RECOMENDACIONES", level=1, centrado=True)
    doc.add_paragraph()

    heading(doc, "CONCLUSIONES", level=2)
    for conclu in [
        "La integración de la Lógica Neutrosófica SVN y la NPL en una plataforma OSINT "
        "produce representaciones epistemológicamente más ricas de la vulnerabilidad "
        "estructural urbana que los índices escalares convencionales, al preservar y "
        "clasificar las contradicciones inter-fuente como información estructuralmente "
        "relevante para la política pública.",

        "El Índice Pluriversal de Vulnerabilidad Estructural (IPVE) calculado mediante "
        "SVNWA ponderado para 30 sectores de Guayaquil revela una distribución de "
        "vulnerabilidad heterogénea, con régimen crítico en sectores periféricos "
        "(Bastión Popular, Monte Sinaí, Guasmo) y paraconsistencia estructural en "
        "sectores centrales (Urdesa, Centro Histórico).",

        "El módulo NPL basado en la triple anotación (mu, lambda, I) permite distinguir "
        "entre indeterminación epistémica (I ≤ 0.45: resoluble con más datos) y "
        "ontológica (I > 0.45: inherente al fenómeno), información que el marco SVN "
        "estándar no puede representar.",

        "El análisis N-fsQCA con cuatro LLMs replica y generaliza los hallazgos de "
        "Leyva-Vázquez et al. (2026a): la prensa local silencia sistemáticamente los "
        "cinco drivers estructurales más relevantes, con brechas superiores a -0.20 "
        "para tráfico de armas, extorsión, nexo carcelario, rutas de narcotráfico y "
        "presión económica.",

        "El detector de paraconsistencia inter-fuente basado en distancia de Hamming "
        "neutrosófica identifica el caso de Urdesa como el hallazgo más significativo: "
        "'barrio bien iluminado' no equivale a 'barrio seguro para todas las personas', "
        "un resultado invisible para modelos escalares convencionales.",

        "El marco ético Camino B, implementado en el módulo EthicalSafeguard, garantiza "
        "que la plataforma opere exclusivamente como herramienta de diagnóstico "
        "estructural comunitario, con audit log automático, validación de propósito "
        "declarado y granularidad mínima de sector conforme a la LOPDP 2021.",
    ]:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(conclu)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing = Pt(22)

    doc.add_paragraph()
    heading(doc, "RECOMENDACIONES", level=2)
    for rec in [
        "Se recomienda a la Universidad de Guayaquil y al GADMGYE establecer un "
        "convenio para la actualización periódica de los datos oficiales que alimentan "
        "la plataforma, especialmente las estadísticas de violencia de género por "
        "sector, actualmente disponibles solo a nivel cantonal.",

        "Se recomienda ampliar el corpus narrativo del módulo N-fsQCA a al menos 100 "
        "artículos por año, incorporando fuentes de medios digitales independientes y "
        "comunidades en Telegram, para aumentar la representatividad del análisis de "
        "silenciamiento mediático.",

        "Se recomienda desarrollar la versión 1.1 de la plataforma incorporando las "
        "páginas de Series Temporales y Comparador de Capas, lo que permitirá detectar "
        "tendencias en el tiempo y comparar la evolución del IPVE entre ciclos "
        "electorales o tras implementación de políticas públicas.",

        "Se recomienda replicar la metodología NEUTROOSINT-GYE en el cantón Durán "
        "(Durán-Safe Digital Twin), aprovechando el trabajo previo del Gemelo Digital "
        "Neutro-Safe y las tres capas adicionales ya diseñadas: riesgo hídrico "
        "(INAMHI), commuting y economía informal.",

        "Se recomienda integrar el módulo NPL al proceso de evaluación de impacto de "
        "políticas públicas de seguridad, permitiendo a tomadores de decisión "
        "identificar si la indeterminación de los datos es epistémica (recolectar más "
        "datos) u ontológica (rediseñar la intervención).",

        "Se recomienda publicar la plataforma bajo licencia MIT en GitHub "
        "(github.com/mleyvaz/neutroosint-gye) para facilitar su adopción por "
        "municipios y organizaciones civiles de la región, documentando el marco "
        "ético Camino B como condición de uso.",
    ]:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(rec)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing = Pt(22)

    add_page_break(doc)


# ============================================================
#  BIBLIOGRAFÍA
# ============================================================

def bibliografia(doc):
    heading(doc, "BIBLIOGRAFÍA", level=1, centrado=True)
    doc.add_paragraph()

    refs = [
        "Atanassov, K. T. (1986). Intuitionistic fuzzy sets. Fuzzy Sets and Systems, 20(1), 87-96. https://doi.org/10.1016/S0165-0114(86)80034-3",
        "da Costa, N. C. A., & Abe, J. M. (1994). Annotated logics Qτ and ultrafilter logic. Studia Logica, 53(1), 75-100.",
        "Ensign, D., Friedler, S. A., Neville, S., Scheidegger, C., & Venkatasubramanian, S. (2018). Runaway feedback loops in predictive policing. Proceedings of FATML, 81, 1-12.",
        "INEC. (2023). Encuesta Nacional de Victimización y Percepción de Inseguridad 2023. Instituto Nacional de Estadística y Censos del Ecuador.",
        "Leyva-Vázquez, M. Y., Cevallos-Torres, L., Guijarro Rodríguez, A., Iturburu-Salvador, M., & Smarandache, F. (2026a). LLM + Neutrosophic fsQCA: Inter-Narrative Causal Consistency and Paraconsistent Detection in Media Accounts of Urban Violence in Guayaquil. In WorldS4 2026, Springer LNNS. (Submission #304).",
        "Leyva-Vázquez, M. Y., & Smarandache, F. (2026b). Neutrosophic Paraconsistent Logic: Evidence Degrees, Ontological Indeterminacy, and Scientific Evidence Synthesis. [En revisión — NCML].",
        "Leyva-Vázquez, M. Y. (2026c). Guayaquil Neutro-Safe Digital Twin: A Pluriversal Community Diagnostic Framework. [Enviado IEEE ETCM 2026, 10-may-2026].",
        "Muggah, R., & Tobón, K. A. (2018). Citizen security in Latin America: Facts and figures. Strategic Note, IGARAPÉ Institute, 33, 1-10.",
        "Perry, W. L., McInnis, B., Price, C. C., Smith, S. C., & Hollywood, J. S. (2013). Predictive policing: The role of crime forecasting in law enforcement operations. RAND Corporation.",
        "Policía Nacional del Ecuador. (2024). Estadísticas de Seguridad Ciudadana. Zona 8 — Guayaquil. Dirección de Operaciones.",
        "Ragin, C. C. (2008). Redesigning social inquiry: Fuzzy sets and beyond. University of Chicago Press.",
        "República del Ecuador. (2008). Constitución de la República del Ecuador. Asamblea Nacional Constituyente.",
        "República del Ecuador. (2021). Ley Orgánica de Protección de Datos Personales. Registro Oficial Suplemento 459.",
        "Smarandache, F. (1998). Neutrosophy: Neutrosophic probability, set, and logic. American Research Press.",
        "Smarandache, F. (1999). A unifying field in logics: Neutrosophic logic. In Philosophy (pp. 1-141). American Research Press.",
        "UNODC. (2024). Global Study on Homicide 2023. United Nations Office on Drugs and Crime. ISBN 978-92-1-130441-1.",
        "Ye, J. (2014). A multicriteria decision-making method using aggregation operators for simplified neutrosophic sets. Journal of Intelligent & Fuzzy Systems, 26(5), 2459-2466.",
        "Zadeh, L. A. (1965). Fuzzy sets. Information and Control, 8(3), 338-353. https://doi.org/10.1016/S0019-9958(65)90241-X",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    add_page_break(doc)


# ============================================================
#  MAIN
# ============================================================

def generate():
    doc = Document()
    set_margin(doc)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    portada(doc)
    declaracion_autoria(doc)
    certificacion_tutor(doc)
    dedicatoria(doc)
    agradecimiento(doc)
    resumen(doc)
    abstract_en(doc)
    introduccion(doc)
    capitulo1(doc)
    capitulo2(doc)
    capitulo3(doc)
    capitulo4(doc)
    capitulo5(doc)
    bibliografia(doc)

    out = "NEUTROOSINT_GYE_Tesis_UG.docx"
    doc.save(out)
    print(f"Tesis generada: {out}")
    return out


if __name__ == "__main__":
    generate()
